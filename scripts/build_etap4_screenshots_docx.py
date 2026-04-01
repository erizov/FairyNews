#!/usr/bin/env python3
"""Word-документ со скриншотами UI (этап 4): главная → шаги 2–3 → отчёты.

Требуется запущенный сервер (например ``uvicorn`` на порту 8765) и
``playwright install chromium``. PNG сохраняются рядом с ``.docx``.

Пример::

    pip install python-docx
    python -m uvicorn app.main:app --host 127.0.0.1 --port 8765
    python scripts/build_etap4_screenshots_docx.py --port 8765
"""

from __future__ import annotations

import argparse
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

_DEFAULT_RUN_ID = "5202d06b-5d38-449a-8ab2-c8335e2b6102"
_NEWS_FILE = ROOT / "docs" / "pipeline_walkthrough_news.txt"


def _wait_health(base: str, timeout: float = 90.0) -> None:
    url = f"{base.rstrip('/')}/api/health"
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        try:
            with urllib.request.urlopen(url, timeout=2) as r:
                if r.status == 200:
                    return
        except (urllib.error.URLError, OSError):
            time.sleep(0.5)
    raise RuntimeError(f"Сервер не отвечает на {url}")


def _news_text() -> str:
    if _NEWS_FILE.is_file():
        return _NEWS_FILE.read_text(encoding="utf-8").strip()
    return (
        "Краткая тестовая новость для скриншота: в тридевятом государстве "
        "починили мост."
    )


def _capture_pngs(base: str, run_id: str, png_dir: Path) -> list[Path]:
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print(
            "Нужен Playwright: pip install playwright && "
            "playwright install chromium",
            file=sys.stderr,
        )
        raise SystemExit(1) from None

    png_dir.mkdir(parents=True, exist_ok=True)
    paths: list[Path] = []
    news = _news_text()

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            viewport={"width": 1100, "height": 900},
            locale="ru-RU",
        )
        page = context.new_page()

        page.goto(f"{base}/", wait_until="networkidle", timeout=120000)
        page.wait_for_selector("#newsList li", timeout=60000)
        page.fill("#customNews", news)
        page.screenshot(path=str(png_dir / "01_main.png"), full_page=True)
        paths.append(png_dir / "01_main.png")

        page.locator("#toStep2").click()
        page.wait_for_selector("#page2.active", timeout=15000)
        page.select_option("#presetSelect", value="russian_folk")
        page.screenshot(path=str(png_dir / "02_step2_preset.png"))
        paths.append(png_dir / "02_step2_preset.png")

        page.locator("#runGen").click()
        page.wait_for_function(
            "() => {"
            "const t = document.getElementById('tale');"
            "return t && t.textContent.length > 50;"
            "}",
            timeout=240000,
        )
        page.screenshot(path=str(png_dir / "03_generated.png"), full_page=True)
        paths.append(png_dir / "03_generated.png")

        page.goto(
            f"{base}/reports-ui/index.html",
            wait_until="networkidle",
            timeout=60000,
        )
        time.sleep(0.5)
        page.screenshot(path=str(png_dir / "04_reports_index.png"))
        paths.append(png_dir / "04_reports_index.png")

        page.goto(
            f"{base}/reports-ui/detail.html?id={run_id}",
            wait_until="networkidle",
            timeout=60000,
        )
        time.sleep(0.4)
        page.screenshot(path=str(png_dir / "05_report_detail.png"), full_page=True)
        paths.append(png_dir / "05_report_detail.png")

        page.goto(
            f"{base}/reports-ui/llm-log.html?id={run_id}",
            wait_until="networkidle",
            timeout=60000,
        )
        time.sleep(0.4)
        page.screenshot(path=str(png_dir / "06_llm_log.png"), full_page=True)
        paths.append(png_dir / "06_llm_log.png")

        browser.close()

    return paths


def _build_docx(
    out_docx: Path,
    png_paths: list[Path],
    titles: list[str],
    descriptions: list[str],
) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        print("pip install python-docx", file=sys.stderr)
        raise SystemExit(1) from None

    doc = Document()
    doc.add_heading("Этап 4 — скриншоты интерфейса FairyNews", level=0)
    doc.add_paragraph(
        "Снимки страниц локального стенда; к каждому — краткое описание."
    )

    width_in = 6.3
    if len(png_paths) != len(titles) or len(titles) != len(descriptions):
        raise ValueError("Несовпадение числа PNG, заголовков и описаний")
    for i, (png, title, desc) in enumerate(
        zip(png_paths, titles, descriptions),
        start=1,
    ):
        if i > 1:
            doc.add_page_break()
        doc.add_heading(f"Страница {i}. {title}", level=1)
        doc.add_paragraph(desc)
        if png.is_file():
            doc.add_picture(str(png), width=Inches(width_in))
        else:
            doc.add_paragraph(f"(файл не найден: {png})")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Скриншоты этапа 4 и сборка Word-документа.",
    )
    parser.add_argument(
        "--host",
        default="127.0.0.1",
        help="хост uvicorn (по умолчанию 127.0.0.1)",
    )
    parser.add_argument(
        "--port",
        type=int,
        default=8765,
        help="порт (по умолчанию 8765)",
    )
    parser.add_argument(
        "--run-id",
        default=_DEFAULT_RUN_ID,
        help="UUID прогона для detail и llm-log",
    )
    parser.add_argument(
        "--out-dir",
        type=Path,
        default=ROOT / "experiments" / "etap4_screenshots",
        help="каталог для PNG и .docx",
    )
    parser.add_argument(
        "--skip-capture",
        action="store_true",
        help="не снимать экран, только собрать .docx из уже лежащих PNG",
    )
    args = parser.parse_args()

    base = f"http://{args.host}:{args.port}"
    out_dir: Path = args.out_dir
    png_dir = out_dir / "png"
    out_docx = out_dir / "etap4_screenshots.docx"

    titles = [
        "Главная",
        "Шаг 2 из 3 — пресет",
        "Сгенерированная сказка (шаг 3)",
        "Сводка отчётов",
        "Детальный отчёт",
        "Журнал вызовов LLM",
    ]
    descriptions = [
        "Стартовая страница: текст новости, при необходимости выбор "
        "карточки из RSS и переход к выбору пресета.",
        "Второй шаг мастера: выбран пресет сказки (например, "
        "«Русский фольклор») перед запуском пайплайна.",
        "Результат генерации: текст сказки, сводка по этапам, вопрос для "
        "проверки понимания и ссылка на сохранённый отчёт.",
        "Таблица сохранённых прогонов: дата, пресет, ссылки на детализацию "
        "и лог LLM.",
        "Развёрнутый отчёт по выбранному прогону: метаданные, RAG, выходы "
        "агентов и эвристики.",
        "Последовательность запросов к модели по тому же прогону "
        "(промпты и ответы в интерфейсе отчёта).",
    ]

    if not args.skip_capture:
        _wait_health(base)
        png_paths = _capture_pngs(base, args.run_id.strip(), png_dir)
    else:
        png_paths = [png_dir / n for n in (
            "01_main.png",
            "02_step2_preset.png",
            "03_generated.png",
            "04_reports_index.png",
            "05_report_detail.png",
            "06_llm_log.png",
        )]

    _build_docx(out_docx, png_paths, titles, descriptions)
    print("Saved:", out_docx)


if __name__ == "__main__":
    main()

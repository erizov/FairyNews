#!/usr/bin/env python3
"""Сборка Word-буклета «Fairy News»: тексты + полностраничные скриншоты UI.

Нужны: работающий сервер (порт по умолчанию 8765), сохранение отчётов
(``FAIRYNEWS_SAVE_REPORTS=1``), ``playwright install chromium``,
``pip install python-docx``.

Сценарий съёмки: главная (полная страница) → шаги 1–3 (пайплайн) → сводка
отчётов → клик по **id** в таблице → детальный отчёт → возврат на сводку →
клик **trace** → журнал LLM. Все снимки ``full_page`` где уместно.

Пример::

    python scripts/build_fairy_news_booklet.py --port 8765
"""

from __future__ import annotations

import argparse
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

_NEWS_FILE = ROOT / "docs" / "pipeline_walkthrough_news.txt"
_OUT_DIR = ROOT / "docs" / "fairy_news_booklet"
_PNG_DIR = _OUT_DIR / "png"


def _wait_health(base: str, timeout: float = 90.0) -> None:
    url = f"{base.rstrip('/')}/api/health"
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        try:
            with urllib.request.urlopen(url, timeout=2) as r:
                if r.status == 200:
                    return
        except (urllib.error.URLError, OSError):
            time.sleep(0.5)
    raise RuntimeError(f"Server not ready: {url}")


def _news_text() -> str:
    if _NEWS_FILE.is_file():
        return _NEWS_FILE.read_text(encoding="utf-8").strip()
    return (
        "A village repaired the old bridge before winter. Officials "
        "checked the plans; neighbors helped carry timber."
    )


def _capture_booklet_pngs(base: str, png_dir: Path) -> list[tuple[Path, str]]:
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print(
            "Install Playwright: pip install playwright && "
            "playwright install chromium",
            file=sys.stderr,
        )
        raise SystemExit(1) from None

    png_dir.mkdir(parents=True, exist_ok=True)
    shots: list[tuple[Path, str]] = []
    news = _news_text()

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            viewport={"width": 1200, "height": 880},
            locale="en-US",
        )
        page = context.new_page()

        def shot(name: str, caption: str, *, full: bool = True) -> None:
            path = png_dir / name
            page.screenshot(path=str(path), full_page=full)
            shots.append((path, caption))

        page.goto(f"{base}/", wait_until="networkidle", timeout=120000)
        page.wait_for_selector("#newsList li", timeout=60000)
        shot("01_home_full.png", "Home — full page (collage, news, step 1).")

        page.fill("#customNews", news)
        shot(
            "02_step1_news_ready.png",
            "Step 1 — custom news text ready; RSS cards available.",
        )

        page.locator("#toStep2").click()
        page.wait_for_selector("#page2.active", timeout=20000)
        page.select_option("#presetSelect", value="russian_folk")
        shot(
            "03_step2_preset.png",
            "Step 2 — tale preset selected (RAG retrieval style).",
            full=False,
        )

        page.locator("#runGen").click()
        page.wait_for_function(
            "() => {"
            "const t = document.getElementById('tale');"
            "return t && t.textContent.length > 50;"
            "}",
            timeout=240000,
        )
        shot(
            "04_step3_pipeline_result.png",
            "Step 3 — generated tale, QA block, links to report & LLM log.",
        )

        time.sleep(1.0)
        page.goto(
            f"{base}/reports-ui/index.html",
            wait_until="networkidle",
            timeout=90000,
        )
        page.wait_for_selector("#tbody tr, #empty", timeout=60000)
        shot(
            "05_reports_index_full.png",
            "Reports index — full page (table with run id and trace).",
        )

        if page.locator("#tbody tr").count() == 0:
            browser.close()
            print(
                "WARN: No saved runs; detail/trace screenshots skipped. "
                "Set FAIRYNEWS_SAVE_REPORTS=1 and re-run.",
                file=sys.stderr,
            )
            return shots

        page.locator("#tbody a[href^='detail.html']").first.click()
        page.wait_for_load_state("networkidle")
        time.sleep(0.5)
        shot(
            "06_detail_after_id_click.png",
            "Detail report — opened by clicking the run id on the index.",
        )

        page.goto(
            f"{base}/reports-ui/index.html",
            wait_until="networkidle",
            timeout=60000,
        )
        page.wait_for_selector("#tbody tr", timeout=30000)
        page.locator("#tbody tr").first.get_by_role("link", name="trace").click()
        page.wait_for_load_state("networkidle")
        time.sleep(0.5)
        shot(
            "07_llm_log_after_trace_click.png",
            "LLM step log — opened by clicking «trace» on the index.",
        )

        browser.close()

    return shots


def _build_docx(
    out_docx: Path,
    shots: list[tuple[Path, str]],
) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError:
        print("pip install python-docx", file=sys.stderr)
        raise SystemExit(1) from None

    doc = Document()
    title = doc.add_heading("Fairy News", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "A booklet for young readers and adults: news seen through a "
        "folk-tale frame — pattern, plot, and connection."
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(12)

    doc.add_page_break()

    doc.add_heading("Why Fairy News?", level=1)
    doc.add_paragraph(
        "Fairy News pairs real headlines with narrative patterns from "
        "folklore and classic tales. The interface keeps a calm, book-like "
        "atmosphere (collage, typography) so reading feels like moving "
        "between today’s news and timeless story shapes."
    )

    doc.add_heading("For young minds", level=2)
    doc.add_paragraph(
        "Children and teens strengthen narrative thinking: they notice "
        "repeating motifs (warning, journey, trial, reward), compare "
        "characters across old and new texts, and practice predicting "
        "what might happen next. Linking a news snippet to a tale "
        "template makes abstract events concrete and memorable."
    )

    doc.add_heading("For adults", level=2)
    doc.add_paragraph(
        "Adults get a reflective lens on the news: archetypes, moral "
        "framing, and distance from pure urgency. The same UI supports "
        "inspection of retrieval (RAG), model steps, and saved reports "
        "for study or demonstration."
    )

    doc.add_heading(
        "Patterns, connections, and foresight",
        level=1,
    )
    doc.add_paragraph(
        "Important skills this tool encourages: (1) connect today’s news "
        "to older stories and shared plots; (2) recognize recurring "
        "characters and roles (helper, trickster, ruler, crowd); "
        "(3) infer what might follow — as in folktales, consequences "
        "often rhyme across centuries. The QA step on the result page "
        "checks understanding and rewards looking ahead."
    )

    doc.add_heading("Generation pipeline (overview)", level=1)
    for line in (
        "1. News input — your text or selected RSS cards (structured slots).",
        "2. Preset — chooses the folklore / corpus angle for similarity search.",
        "3. Agents — news structuring → tale draft → audit → Q&A.",
        "4. RAG — retrieves tale fragments to ground the story.",
        "5. Report + LLM log — full JSON report and per-step prompts/responses.",
    ):
        doc.add_paragraph(line, style="List Number")

    doc.add_page_break()
    doc.add_heading("Screenshots from the live app", level=1)
    doc.add_paragraph(
        "The following pages show the running site (example: "
        "http://127.0.0.1:8765/) and reports UI, including navigation "
        "via the run id and the trace link as a student or reviewer would."
    )

    width = Inches(6.35)
    for path, caption in shots:
        doc.add_page_break()
        cap_p = doc.add_paragraph()
        cap_run = cap_p.add_run(caption)
        cap_run.bold = True
        if path.is_file():
            doc.add_picture(str(path), width=width)
        else:
            doc.add_paragraph(f"(missing image: {path.name})")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def main() -> None:
    parser = argparse.ArgumentParser(description="Build Fairy News booklet DOCX.")
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8765)
    parser.add_argument(
        "--skip-capture",
        action="store_true",
        help="only build DOCX from existing PNGs in docs/fairy_news_booklet/png",
    )
    args = parser.parse_args()

    base = f"http://{args.host}:{args.port}"
    out_docx = _OUT_DIR / "Fairy_News_Booklet.docx"

    if args.skip_capture:
        order = [
            ("01_home_full.png", "Home — full page."),
            ("02_step1_news_ready.png", "Step 1 — news ready."),
            ("03_step2_preset.png", "Step 2 — preset."),
            (
                "04_step3_pipeline_result.png",
                "Step 3 — pipeline result.",
            ),
            ("05_reports_index_full.png", "Reports index."),
            (
                "06_detail_after_id_click.png",
                "Detail via id click.",
            ),
            (
                "07_llm_log_after_trace_click.png",
                "LLM log via trace click.",
            ),
        ]
        shots = [(_PNG_DIR / n, c) for n, c in order]
    else:
        _wait_health(base)
        shots = _capture_booklet_pngs(base, _PNG_DIR)

    _build_docx(out_docx, shots)
    print("Saved:", out_docx)


if __name__ == "__main__":
    main()

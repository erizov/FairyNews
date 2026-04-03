#!/usr/bin/env python3
"""Русский буклет Fairy News: тексты + содержимое страниц из DOM (без скриншотов).

Playwright сохраняет **полный HTML** в ``html/``. Основной буклет —
«продающий» текст для школы, родчатов и Habr **без** длинных DOM-вставок;
полные выгрузки — в **отдельном** ``…_Prilozhenie_DOM.docx``. Есть короткая
**Lite**-версия для соцсетей.

Демо-URL для буклета: переменная ``FAIRYNEWS_BOOKLET_DEMO_URL`` (иначе
ссылка на репозиторий).

Нужны: сервер, ``FAIRYNEWS_SAVE_REPORTS=1``, playwright, python-docx,
beautifulsoup4.

Пример::

    python scripts/build_fairy_news_ru_booklet.py --port 8765
"""

from __future__ import annotations

import argparse
import hashlib
import os
import sys
import time
import urllib.error
import urllib.request
from pathlib import Path
from typing import Any

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

_NEWS_FILE = ROOT / "docs" / "pipeline_walkthrough_news.txt"
_OUT_DIR = ROOT / "docs" / "fairy_news_booklet_ru"
_HTML_DIR = _OUT_DIR / "html"
_MAIN_NAME = "Fairy_News_RU_Booklet.docx"
_LITE_NAME = "Fairy_News_RU_Booklet_Lite.docx"
_APPENDIX_NAME = "Fairy_News_RU_Booklet_Prilozhenie_DOM.docx"

# Те же файлы, что в ``scripts/download_collage_images.py`` (Commons).
_LITE_COLLAGE_WIKI: dict[str, str] = {
    "tile-00": "Newspaper_stack.jpg",
    "tile-01": "Old_paper.jpg",
    "tile-02": "London_from_a_hot_air_balloon.jpg",
    "tile-03": "Pleiades_large.jpg",
    "tile-04": "Open_book.jpg",
    "tile-05": (
        "Two_bookshelves_full_of_books_belonging_to_Unitedmissionary_"
        "(2010).jpg"
    ),
    "tile-06": "Forest_path.jpg",
    "tile-07": "Wood_texture.jpg",
}
_LITE_IMAGE_UA = (
    "FairyNewsBooklet/1.0 "
    "(https://github.com/erizov/FairyNews; RU Lite booklet image)"
)


def _wikimedia_commons_url(filename: str) -> str:
    digest = hashlib.md5(filename.encode("utf-8")).hexdigest()
    return (
        "https://upload.wikimedia.org/wikipedia/commons/"
        f"{digest[0]}/{digest[0:2]}/{filename}"
    )


def _lite_tile_raster_bytes(stem: str) -> bytes | None:
    """Растр из ``frontend/collage`` или загрузка с Wikimedia Commons."""
    from app.api_schemas import _collage_assets_dir, _resolved_collage_file

    filename, _ = _resolved_collage_file(stem)
    base = _collage_assets_dir()
    path = base / filename
    raster_ext = (".webp", ".png", ".jpg", ".jpeg")
    if path.is_file() and path.suffix.lower() in raster_ext:
        return path.read_bytes()

    wiki_name = _LITE_COLLAGE_WIKI.get(stem)
    if not wiki_name:
        return None
    url = _wikimedia_commons_url(wiki_name)
    req = urllib.request.Request(
        url,
        headers={
            "User-Agent": _LITE_IMAGE_UA,
            "Accept": "image/*,*/*;q=0.8",
        },
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            return resp.read()
    except OSError:
        return None


def _add_lite_internet_image(doc: Any, stem: str) -> None:
    """Одна иллюстрация после блока (Lite); источник — коллаж / Commons."""
    from io import BytesIO

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    data = _lite_tile_raster_bytes(stem)
    if not data:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Inches(0.08)
    p.paragraph_format.space_after = Inches(0.12)
    run = p.add_run()
    run.add_picture(BytesIO(data), width=Inches(4.0))


def _demo_url() -> str:
    u = os.environ.get("FAIRYNEWS_BOOKLET_DEMO_URL", "").strip()
    if u:
        return u
    return "https://github.com/erizov/FairyNews"


def _wait_health(base: str, timeout: float = 90.0) -> None:
    url = f"{base.rstrip('/')}/api/health"
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        try:
            with urllib.request.urlopen(url, timeout=2) as r:
                if r.status == 200:
                    return
        except (urllib.error.URLError, OSError):
            time.sleep(0.5)
    raise RuntimeError(f"Сервер не отвечает: {url}")


def _news_text() -> str:
    if _NEWS_FILE.is_file():
        return _NEWS_FILE.read_text(encoding="utf-8").strip()
    return (
        "В тридевятом государстве починили мост перед зимой. "
        "Соседи несли брёвна, писари сверяли сроки."
    )


def _save_html(html_dir: Path, name: str, page: Any) -> Path:
    html_dir.mkdir(parents=True, exist_ok=True)
    path = html_dir / name
    content = page.content()
    path.write_text(content, encoding="utf-8")
    return path


def _capture_flow(base: str, html_dir: Path) -> list[tuple[Path, str]]:
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        print(
            "Установите Playwright: pip install playwright && "
            "playwright install chromium",
            file=sys.stderr,
        )
        raise SystemExit(1) from None

    chunks: list[tuple[Path, str]] = []
    news = _news_text()

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        context = browser.new_context(
            viewport={"width": 1200, "height": 880},
            locale="ru-RU",
        )
        page = context.new_page()

        def snap(fname: str, title_ru: str) -> None:
            path = _save_html(html_dir, fname, page)
            chunks.append((path, title_ru))

        page.goto(f"{base}/", wait_until="networkidle", timeout=120000)
        page.wait_for_selector("#newsList li", timeout=60000)
        snap("01_glavnaya.html", "Главная — полный HTML (коллаж, шаг 1)")

        page.fill("#customNews", news)
        snap(
            "02_shag1_novost.html",
            "Шаг 1 — введён текст новости, список RSS в DOM",
        )

        page.locator("#toStep2").click()
        page.wait_for_selector("#page2.active", timeout=20000)
        page.select_option("#presetSelect", value="russian_folk")
        snap("03_shag2_preset.html", "Шаг 2 — выбран пресет RAG")

        page.locator("#runGen").click()
        page.wait_for_function(
            "() => {"
            "const t = document.getElementById('tale');"
            "return t && t.textContent.length > 50;"
            "}",
            timeout=240000,
        )
        snap(
            "04_shag3_rezultat.html",
            "Шаг 3 — сказка, вопрос, ссылки на отчёт и LLM",
        )

        time.sleep(1.0)
        page.goto(
            f"{base}/reports-ui/index.html",
            wait_until="networkidle",
            timeout=90000,
        )
        page.wait_for_selector("#tbody tr, #empty", timeout=60000)
        snap(
            "05_otchety_index.html",
            "Сводка отчётов — таблица (id, trace) в DOM",
        )

        if page.locator("#tbody tr").count() == 0:
            browser.close()
            print(
                "ПРЕДУПРЕЖДЕНИЕ: нет сохранённых прогонов — деталка и LLM "
                "пропущены. Задайте FAIRYNEWS_SAVE_REPORTS=1.",
                file=sys.stderr,
            )
            return chunks

        page.locator("#tbody a[href^='detail.html']").first.click()
        page.wait_for_load_state("networkidle")
        time.sleep(0.6)
        snap(
            "06_detal_posle_klik_id.html",
            "Детальный отчёт — после клика по id в таблице",
        )

        page.goto(
            f"{base}/reports-ui/index.html",
            wait_until="networkidle",
            timeout=60000,
        )
        page.wait_for_selector("#tbody tr", timeout=30000)
        page.locator("#tbody tr").first.get_by_role("link", name="trace").click()
        page.wait_for_load_state("networkidle")
        time.sleep(0.6)
        snap(
            "07_llm_log_posle_trace.html",
            "Журнал LLM — после клика по «trace»",
        )

        browser.close()

    return chunks


def _html_table_to_doc(doc: Any, table_el: Any) -> None:
    rows_el = table_el.find_all("tr", recursive=False)
    if not rows_el:
        rows_el = table_el.find_all("tr")
    if not rows_el:
        return
    max_cols = 0
    parsed: list[list[str]] = []
    for tr in rows_el:
        cells = tr.find_all(["th", "td"], recursive=False)
        row_t = [c.get_text("\n", strip=True) for c in cells]
        max_cols = max(max_cols, len(row_t))
        parsed.append(row_t)
    if max_cols == 0:
        return
    tbl = doc.add_table(rows=len(parsed), cols=max_cols)
    tbl.style = "Table Grid"
    for i, row_t in enumerate(parsed):
        for j in range(max_cols):
            txt = row_t[j] if j < len(row_t) else ""
            tbl.rows[i].cells[j].text = txt


def _walk_html_to_doc(doc: Any, root: Any) -> None:
    from bs4 import NavigableString, Tag
    from docx.shared import Pt

    skip_tags = frozenset({"script", "style", "noscript"})
    mono = Pt(9)

    def walk(el: Any) -> None:
        if not isinstance(el, Tag):
            return
        if el.name in skip_tags:
            return
        if el.get("aria-hidden") == "true":
            return
        if el.name == "table":
            _html_table_to_doc(doc, el)
            return
        if el.name in ("h1", "h2", "h3", "h4"):
            lvl = int(el.name[1])
            t = el.get_text(" ", strip=True)
            if t:
                doc.add_heading(t, level=min(lvl, 3))
            return
        if el.name == "p":
            t = el.get_text(" ", strip=True)
            if t:
                doc.add_paragraph(t)
            return
        if el.name == "pre":
            t = el.get_text()
            if t.strip():
                p = doc.add_paragraph()
                r = p.add_run(t.rstrip())
                r.font.name = "Courier New"
                r.font.size = mono
            return
        if el.name == "textarea":
            t = el.get_text()
            if t.strip():
                doc.add_paragraph("Текст поля (textarea):")
                p = doc.add_paragraph()
                r = p.add_run(t.strip())
                r.font.name = "Courier New"
                r.font.size = mono
            return
        if el.name in ("ul", "ol"):
            for li in el.find_all("li", recursive=False):
                lt = li.get_text(" ", strip=True)
                if lt:
                    doc.add_paragraph(lt, style="List Bullet")
            return
        for child in el.children:
            if isinstance(child, NavigableString):
                tx = str(child).strip()
                if tx:
                    doc.add_paragraph(tx)
            else:
                walk(child)

    walk(root)


def _append_html_snapshot(doc: Any, html_path: Path, section_title: str) -> None:
    from bs4 import BeautifulSoup

    doc.add_page_break()
    doc.add_heading(section_title, level=1)
    note = doc.add_paragraph(
        f"Источник DOM (файл): {html_path.name} — полная разметка "
        "страницы после загрузки; ниже извлечённый текст и таблицы."
    )
    if note.runs:
        note.runs[0].italic = True

    raw = html_path.read_text(encoding="utf-8")
    soup = BeautifulSoup(raw, "html.parser")
    root = (
        soup.find("main")
        or soup.find(id="main")
        or soup.find("div", class_="wrap")
        or soup.body
    )
    if root is None:
        root = soup
    _walk_html_to_doc(doc, root)


def _add_collage_grid_first_page(doc: Any) -> None:
    """Вставляет текущий коллаж (те же файлы, что у ``/static/collage``)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    from app.api_schemas import _COLLAGE_TILES, _collage_assets_dir, _resolved_collage_file

    pic_ext = frozenset({".jpg", ".jpeg", ".png", ".webp"})
    cap = doc.add_paragraph(
        "Коллаж интерфейса — те же изображения, что на сайте:"
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(10)
        run.italic = True

    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = "Table Grid"
    base_dir = _collage_assets_dir()
    cell_w = Inches(1.48)

    for idx, (stem, alt_ru, _motif) in enumerate(_COLLAGE_TILES):
        filename, _fb = _resolved_collage_file(stem)
        path = base_dir / filename
        row_i, col_i = divmod(idx, 4)
        cell = tbl.rows[row_i].cells[col_i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.is_file() and path.suffix.lower() in pic_ext:
            run = p.add_run()
            run.add_picture(str(path), width=cell_w)
        else:
            p.add_run(f"[{filename}]")

    alts = [a for _stem, a, _m in _COLLAGE_TILES]
    leg = doc.add_paragraph("Кадры (подписи): " + " · ".join(alts))
    for lr in leg.runs:
        lr.font.size = Pt(8)
    leg.paragraph_format.space_before = Pt(6)


def _add_cover_and_epigraph(doc: Any, *, lite: bool) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    t = doc.add_heading("Fairy News", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(
        "Как читать новости через сказку: спокойнее, нагляднее, "
        "с пользой для класса и семьи"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(13)
        run.bold = True
    aud = doc.add_paragraph("Школа · родительские чаты · Хабр")
    aud.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in aud.runs:
        run.font.size = Pt(11)
        run.italic = True

    doc.add_paragraph()
    ep = doc.add_paragraph(
        "«Сказка ложь, да в ней намёк — добрым молодцам урок.»"
    )
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in ep.runs:
        run.italic = True
        run.font.size = Pt(10)
    if not lite:
        doc.add_paragraph(
            "Намёк здесь простой: в узнаваемом сюжете легче держать "
            "мысль, чем в сухом заголовке ленты."
        )


def _add_why_me_three_bullets(doc: Any) -> None:
    doc.add_heading("Зачем мне это — в трёх строках", level=1)
    doc.add_paragraph(
        "Ребёнку и подростку — тренировка узоров сказки и вопроса "
        "«что будет дальше?» без страха ошибиться.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Взрослому — дистанция, архетипы и меньше тревожного "
        "«рваного» скролла; факты по-прежнему проверяются вне сказки.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Семье за столом — общий язык: «это как в сказке про…», "
        "спор мягче, память дольше.",
        style="List Bullet",
    )


def _add_channels_school_parents_habr(doc: Any, *, lite: bool) -> None:
    doc.add_heading("Где это заходит лучше всего", level=1)

    doc.add_heading("Школа", level=2)
    doc.add_paragraph(
        "10–20 минут на уроке литературы или обществознания: одна "
        "нейтральная новость → пресет «народная сказка» → обсуждение "
        "сюжетного созвучия. Не нужно ставить сервер на каждый ПК: "
        "достаточно демонстрации с проектора или заранее сохранённого "
        "отчёта из приложения."
    )
    if not lite:
        doc.add_paragraph(
            "Сопоставьте с программой: устный народный творчество, "
            "жанры фольклора, медиаграмотность — Fairy News даёт живой "
            "материал для анализа, а не заменяет учебник."
        )
    if lite:
        _add_lite_internet_image(doc, "tile-02")

    doc.add_heading("Родительские чаты", level=2)
    doc.add_paragraph(
        "Короткий пост из готового блока ниже + ссылка на демо. "
        "Акцент: не «ещё один бот», а совместное чтение и один "
        "вопрос ребёнку после сказки."
    )
    if not lite:
        doc.add_paragraph(
            "В чате лучше не спорить о политике на свежей ленте — "
            "взять заранее выбранный или свой спокойный текст новости, "
            "чтобы тон остался дружелюбным."
        )
    if lite:
        _add_lite_internet_image(doc, "tile-06")

    doc.add_heading("Хабр и техническое сообщество", level=2)
    doc.add_paragraph(
        "Стек: FastAPI, RAG по снимку корпуса / Chroma, мультиагентный "
        "пайплайн (новость → сказка → аудит → QA), сохранение отчётов в "
        "SQLite, UI на статическом фронте. Прозрачность: таблица прогонов, "
        "детальный JSON-отчёт, пошаговый журнал запросов к LLM (trace)."
    )
    if not lite:
        doc.add_paragraph(
            "Репозиторий с инструкциями и скриптами экспериментов — "
            "см. ссылку на демо ниже; воспроизводимость важнее «магии»."
        )
    if lite:
        _add_lite_internet_image(doc, "tile-05")


def _add_mini_case(doc: Any) -> None:
    doc.add_heading("Мини-кейс: одна новость — три угла", level=1)
    doc.add_paragraph(
        "Условная новость: «В посёлке до зимы отремонтировали мост, "
        "жители помогали с материалами»."
    )
    doc.add_paragraph(
        "Угол сказки: узнаваемый мотив общего дела, срока и награды "
        "за труд — мост как «чудо» коллективными усилиями."
    )
    doc.add_paragraph(
        "Вопрос ребёнку: «Кто в этой истории похож на сказочного "
        "помощника? Что могло пойти не так, как в сказках про лень?»"
    )
    doc.add_paragraph(
        "Тема для взрослого: ответственность сообщества и прозрачность "
        "решений — без морализаторства, через узнаваемый сюжет."
    )


def _add_glossary(doc: Any, *, lite: bool) -> None:
    doc.add_heading("Глоссарий в одном предложении", level=1)
    rows = [
        ("Пресет", "Заранее заданный угол поиска по корпусу сказок."),
        ("RAG", "Поиск релевантных фрагментов текстов и подстановка в промпт."),
        ("Агент (шаг LLM)", "Отдельный вызов модели: новость, сказка, аудит, QA."),
        ("Отчёт прогона", "Сохранённый JSON: новость, сказка, RAG, эвристики."),
        ("Trace / журнал LLM", "Промпты и ответы по каждому шагу для разбора."),
        ("RSS-карточки", "Короткие свежие заголовки из лент; можно не брать."),
    ]
    if lite:
        rows = rows[:3]
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    for i, (k, v) in enumerate(rows):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v


def _add_ethics_ai(doc: Any) -> None:
    doc.add_heading("Честность и границы", level=1)
    doc.add_paragraph(
        "Это не гадание и не предсказание судьбы: «прогноз» в буклете — "
        "учебная игра сюжетными шаблонами, как в разборе литературы."
    )
    doc.add_paragraph(
        "Иногда ошибается и человек, и модель: сказка может исказить "
        "факты новости. Всегда сверяйтесь с первоисточниками и здравым "
        "смыслом; отчёт и trace показывают, на чём стоял ответ."
    )


def _add_pipeline_diagram(doc: Any) -> None:
    doc.add_heading("Пайплайн — схема", level=1)
    diagram = (
        "  [Новость или RSS]     [Пресет корпуса]\n"
        "         \\                    /\n"
        "          v                  v\n"
        "       [Поиск RAG — фрагменты сказок]\n"
        "                  |\n"
        "                  v\n"
        "    [Агент: структура новости]\n"
        "                  |\n"
        "                  v\n"
        "    [Агент: черновик сказки]\n"
        "                  |\n"
        "         +--------+--------+\n"
        "         v                 v\n"
        "   [Аудит сказки]    [Вопрос и эталон]\n"
        "         \\                 /\n"
        "          v               v\n"
        "     [Сохранённый отчёт]  [Журнал LLM]"
    )
    p = doc.add_paragraph()
    r = p.add_run(diagram)
    r.font.name = "Courier New"
    doc.add_heading("Пайплайн — список шагов", level=2)
    for line in (
        "1. Ввод новости — свой текст или отмеченные карточки RSS.",
        "2. Пресет — угол поиска по корпусу сказок (RAG).",
        "3. Агенты — структурирование новости → черновик сказки → аудит "
        "→ вопрос–ответ.",
        "4. RAG — фрагменты сказок для опоры текста.",
        "5. Отчёт и журнал LLM — JSON и промпты/ответы по шагам.",
    ):
        doc.add_paragraph(line, style="List Number")


def _add_checklist_and_post(doc: Any, *, lite: bool) -> None:
    doc.add_heading("Чек-лист: 5 вопросов после сказки", level=1)
    qs = (
        "О какой старой сказке или былине это напомнило?",
        "Кто здесь «помогает», кто «мешает», кто «решает»?",
        "Что может произойти дальше, если сюжет как в народной сказке?",
        "Что в новости осталось важным, даже если сказка упростила?",
        "Какой один факт из новости стоит проверить по источнику?",
    )
    for q in qs:
        doc.add_paragraph(q, style="List Number")

    if lite:
        _add_lite_internet_image(doc, "tile-06")

    doc.add_heading("Готовый пост для ВК / Telegram", level=1)
    doc.add_paragraph(
        "Попробуйте Fairy News: одна спокойная новость превращается в "
        "короткую сказку в стиле русского фольклора — удобно обсуждать "
        "с детьми без политического накала ленты."
    )
    doc.add_paragraph(
        "После генерации задайте ребёнку один вопрос со страницы результата "
        "— это тренирует внимание и воображение, а не «заменяет» чтение "
        "новостей."
    )
    doc.add_paragraph(
        "Ссылка на проект и демо — в комментарии; исходники открыты, "
        "есть отчёты и журнал шагов модели для любопытных."
    )

    if lite:
        _add_lite_internet_image(doc, "tile-00")


def _add_lite_defense_section(doc: Any) -> None:
    """Требования к защите и привязка к проекту (только Lite-буклет)."""
    doc.add_heading("Презентация на защите диплома", level=1)
    doc.add_paragraph(
        "Этап 5 (GPT Engineer): подготовка к публичному выступлению при защите "
        "дипломного проекта — очно или дистанционно, регламент до 20 минут. "
        "Нужно в свободной форме представить исследование и дополнить рассказ "
        "визуально, слайдами."
    )
    doc.add_paragraph(
        "В презентации нужно отразить всё, что вы планируете рассказать о "
        "своём проекте: цели, задачи, базу знаний, её сбор и обработку, "
        "блок-схему, алгоритм, эксперименты, трудности и их решение, "
        "результат, выводы, заключение."
    )
    doc.add_paragraph(
        "На защите слушатель делится экраном с презентацией и слайд за "
        "слайдом рассказывает о проекте и работе над ним."
    )
    doc.add_paragraph(
        "На защите вас могут попросить продемонстрировать работоспособность "
        "алгоритма: позаботьтесь, чтобы в этот момент код был запущен — "
        "поднят сервер, открыт интерфейс, доступна модель (GigaChat и т.п.), "
        "можно пройти путь от новости до сказки и сохранённого отчёта."
    )

    doc.add_heading("Ориентир для слайдов: Fairy News", level=2)
    doc.add_paragraph(
        "Цели — снизить сухую тревогу ленты и развить нарративное "
        "мышление за счёт сюжетных параллелей с фольклором; задачи — "
        "реализовать веб-сервис с RAG по корпусу сказок, мультиагентным "
        "пайплайном LLM и прозрачными отчётами.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "База знаний — текстовый корпус народных сказок (снимок для RAG, "
        "индексация в векторном хранилище); сбор и обработка — загрузка "
        "корпуса, нарезка на фрагменты, пресеты поиска, подстановка в "
        "промпты.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Блок-схема и алгоритм — слайд с общей схемой (новость → RAG → цепочка "
        "агентов: структура, черновик сказки, аудит, вопрос–ответ) и при "
        "желании отдельный слайд по шагам и данным (см. раздел «Пайплайн» "
        "в этом буклете).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Эксперименты — сравнение пресетов и промптов, прогоны из ноутбуков "
        "и каталога experiments, разбор отчётов и журнала LLM в UI.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Трудности и решения — перечислены в подразделе ниже (RAG, связь с "
        "LLM, выбор модели, мультиагентный аудит, локальные скрипты запуска).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Результат, выводы, заключение — работающий прототип, воспроизводимый "
        "код, польза для обучения и демонстрации; выводы о роли RAG и сюжетных "
        "шаблонов; заключение о достигнутой полноте целей и направлениях "
        "развития.",
        style="List Bullet",
    )

    doc.add_heading("Проблемы и решения (по проекту)", level=2)
    doc.add_paragraph(
        "RAG по корпусу сказок не закрывает весь нужный контекст. Решение — "
        "локальная подгрузка материалов из текстовых файлов, PDF и других "
        "форматов (в корпус или в промпт: пресеты, пояснения на слайдах; не "
        "только выдача поиска по индексу).",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Стабильная связь с провайдером LLM. Решение — smoke-тест "
        "``app/llm_connect_try.py``: загрузка переменных из ``.env``, "
        "создание OpenAI-совместимого клиента (AITunnel и аналоги), один "
        "короткий запрос к чату — тот же стек, что у основного пайплайна и "
        "API, без «угадывания» ключа и базового URL на защите.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Выбор рабочей модели. Решение — не интуитивно, а по результатам "
        "тестов: качество черновика и аудита, длина ответа, задержка, устойчивость "
        "к ошибкам API; зафиксировать вывод на слайде «эксперименты».",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Риск «голого» черновика без проверки. Решение — мультиагентная "
        "цепочка: отдельный шаг аудита оценивает целостность и релевантность "
        "результата по отношению к новости и к фрагментам RAG перед выдачей "
        "пользователю.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Предсказуемый локальный запуск для демонстрации. Решение — скрипты в "
        "``scripts/``: ``start_web.ps1`` и ``start_web.sh``, "
        "``stop_web.ps1`` и ``stop_web.sh``, ``restart_web.ps1`` и "
        "``restart_web.sh`` — старт, остановка и перезапуск веб-сервиса без "
        "ручного поиска процесса и порта.",
        style="List Bullet",
    )


def _add_appendix_pointer(
    doc: Any,
    *,
    appendix_name: str,
    demo: str,
) -> None:
    doc.add_page_break()
    doc.add_heading("Приложение и материалы", level=1)
    doc.add_paragraph(
        f"Полные выгрузки текста из интерфейса (все шаги и длинные поля) "
        f"вынесены в отдельный файл вложения:"
    )
    p = doc.add_paragraph()
    r = p.add_run(appendix_name)
    r.bold = True
    doc.add_paragraph(
        f"Исходные снимки DOM (HTML): каталог "
        f"``{_HTML_DIR.relative_to(ROOT)}`` — можно открыть в браузере."
    )
    doc.add_paragraph("Демо и репозиторий (подставьте QR при печати):")
    doc.add_paragraph(demo)
    doc.add_paragraph(
        "Переменная окружения FAIRYNEWS_BOOKLET_DEMO_URL задаёт URL для "
        "этой строки при следующей сборке буклета."
    )


def _build_marketing_docx(
    out_docx: Path,
    *,
    lite: bool,
    appendix_name: str,
) -> None:
    try:
        from docx import Document
    except ImportError:
        print("pip install python-docx", file=sys.stderr)
        raise SystemExit(1) from None

    doc = Document()
    _add_cover_and_epigraph(doc, lite=lite)
    if lite:
        _add_lite_internet_image(doc, "tile-03")
    _add_collage_grid_first_page(doc)
    if lite:
        _add_lite_internet_image(doc, "tile-01")

    doc.add_page_break()
    _add_why_me_three_bullets(doc)
    if lite:
        _add_lite_internet_image(doc, "tile-00")

    doc.add_page_break()
    doc.add_heading("Зачем Fairy News", level=1)
    doc.add_paragraph(
        "Сервис сопоставляет актуальные формулировки с устойчивыми "
        "сюжетными ходами фольклора и классических сказок. Интерфейс "
        "сохраняет спокойную «книжную» атмосферу: коллаж, типографика — "
        "чтобы чтение напоминало переход от сегодняшней хроники к "
        "проверенным временем формам рассказа."
    )
    if lite:
        _add_lite_internet_image(doc, "tile-04")
    if not lite:
        doc.add_heading("Для детей и подростков", level=2)
        doc.add_paragraph(
            "Развивается нарративное мышление: замечаются повторяющиеся "
            "мотивы (предупреждение, дорога, испытание, награда), "
            "сравниваются персонажи старых и новых текстов, тренируется "
            "предсказание продолжения."
        )
        doc.add_heading("Для взрослых", level=2)
        doc.add_paragraph(
            "Дистанция и опора на архетипы; тот же интерфейс позволяет "
            "разбирать RAG, шаги модели и сохранённые отчёты — для учёбы "
            "и демонстрации."
        )

    doc.add_page_break()
    _add_channels_school_parents_habr(doc, lite=lite)

    doc.add_page_break()
    _add_mini_case(doc)
    if lite:
        _add_lite_internet_image(doc, "tile-07")

    doc.add_page_break()
    _add_glossary(doc, lite=lite)
    if lite:
        _add_lite_internet_image(doc, "tile-05")

    doc.add_page_break()
    _add_ethics_ai(doc)
    if lite:
        _add_lite_internet_image(doc, "tile-03")

    doc.add_page_break()
    doc.add_heading("Связи, сюжеты, прогноз", level=1)
    doc.add_paragraph(
        "Важно научиться: (1) видеть связь между новостью и старыми "
        "историями; (2) узнавать повторяющиеся роли; (3) предполагать "
        "возможное продолжение. Блок вопроса на шаге 3 проверяет понимание."
    )
    if lite:
        _add_lite_internet_image(doc, "tile-02")
    if not lite:
        doc.add_paragraph(
            "Интертекстуальность помогает не застревать в одном заголовке: "
            "та же тема встречалась в былинах и в сегодняшних лентах."
        )
        doc.add_paragraph(
            "Сюжетные ходы работают как подсказки для осторожного прогноза "
            "и внимания к источникам."
        )
        doc.add_paragraph(
            "Прогноз — не гадание, а осмысленное ожидание по узнаваемым "
            "ролям и мотивам."
        )
        doc.add_paragraph(
            "Сказочная рамка смягчает сухую тревогу ленты при сохранении "
            "серьёзности темы."
        )

    doc.add_page_break()
    doc.add_heading(
        "Спокойствие и учёба на примерах — в любом возрасте",
        level=1,
    )
    doc.add_paragraph(
        "Fairy News не похож на экзамен: можно читать медленно и "
        "возвращаться к пресету. Обучение на примерах объединяет "
        "поколения; ошибка в прогнозе здесь учит, а не наказывает."
    )
    if lite:
        _add_lite_internet_image(doc, "tile-01")
    if not lite:
        doc.add_paragraph(
            "Примеры из корпуса — безопасная площадка: открыть отчёт и "
            "увидеть, как модель связала новость с фольклором."
        )

    doc.add_page_break()
    _add_pipeline_diagram(doc)
    if lite:
        _add_lite_internet_image(doc, "tile-00")

    doc.add_page_break()
    _add_checklist_and_post(doc, lite=lite)

    if lite:
        doc.add_page_break()
        _add_lite_defense_section(doc)
        _add_lite_internet_image(doc, "tile-04")

    demo = _demo_url()
    _add_appendix_pointer(doc, appendix_name=appendix_name, demo=demo)
    if lite:
        _add_lite_internet_image(doc, "tile-07")

    if lite:
        doc.add_paragraph()
        tail = doc.add_paragraph()
        tr = tail.add_run(
            "Полная версия буклета (больше текста и разделов): "
            f"{_MAIN_NAME}"
        )
        tr.italic = True

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def _build_appendix_docx(
    out_docx: Path,
    html_chunks: list[tuple[Path, str]],
) -> None:
    try:
        from docx import Document
    except ImportError:
        print("pip install python-docx", file=sys.stderr)
        raise SystemExit(1) from None

    doc = Document()
    doc.add_heading("Приложение: выгрузка DOM в текст", level=0)
    doc.add_paragraph(
        "Этот файл — вложение к основному буклету "
        f"({_MAIN_NAME}). Здесь извлечённое содержимое страниц после "
        "прогона в браузере (сказка, отчёты, JSON журнала LLM). "
        f"Исходные HTML: ``{_HTML_DIR.relative_to(ROOT)}``."
    )

    for path, title in html_chunks:
        if path.is_file():
            _append_html_snapshot(doc, path, title)
        else:
            doc.add_page_break()
            doc.add_paragraph(f"(нет файла: {path.name})")

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Сборка русских буклетов Fairy News (маркетинг + приложение).",
    )
    parser.add_argument("--host", default="127.0.0.1")
    parser.add_argument("--port", type=int, default=8765)
    parser.add_argument(
        "--skip-capture",
        action="store_true",
        help="только DOCX из html в docs/fairy_news_booklet_ru/html",
    )
    args = parser.parse_args()

    base = f"http://{args.host}:{args.port}"
    out_main = _OUT_DIR / _MAIN_NAME
    out_lite = _OUT_DIR / _LITE_NAME
    out_app = _OUT_DIR / _APPENDIX_NAME

    if args.skip_capture:
        order = [
            ("01_glavnaya.html", "Главная — полный HTML"),
            ("02_shag1_novost.html", "Шаг 1 — новость"),
            ("03_shag2_preset.html", "Шаг 2 — пресет"),
            ("04_shag3_rezultat.html", "Шаг 3 — результат"),
            ("05_otchety_index.html", "Сводка отчётов"),
            ("06_detal_posle_klik_id.html", "Детальный отчёт"),
            ("07_llm_log_posle_trace.html", "Журнал LLM"),
        ]
        html_chunks = [(_HTML_DIR / n, t) for n, t in order]
    else:
        _wait_health(base)
        html_chunks = _capture_flow(base, _HTML_DIR)

    _build_marketing_docx(
        out_main,
        lite=False,
        appendix_name=_APPENDIX_NAME,
    )
    _build_marketing_docx(
        out_lite,
        lite=True,
        appendix_name=_APPENDIX_NAME,
    )
    _build_appendix_docx(out_app, html_chunks)
    print("Saved:", out_main)
    print("Saved:", out_lite)
    print("Saved:", out_app)


if __name__ == "__main__":
    main()
